import logging
import os
import io
from telegram import Update, ReplyKeyboardMarkup, ReplyKeyboardRemove, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, MessageHandler, CallbackQueryHandler, ContextTypes, filters
from PIL import Image, ImageEnhance, ImageFilter
from docx import Document
from docx.shared import Inches, Pt
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.utils import ImageReader

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

BOT_TOKEN = os.environ.get("BOT_TOKEN")
ADMIN_ID = 8016799878

# Fayl saqlash
# {"darsliklar": {"1-sinf": {"Matematika 1-qism": file_id}}}
# {"badiiy": {"Roman": {"O'tkan kunlar": file_id}}}
# {"pedagogik": {"Dars ishlanmalar": {"Matematika 5-sinf": file_id}}}
file_storage = {
    "darsliklar": {},
    "badiiy": {},
    "pedagogik": {}
}

DARSLIKLAR = {
    "1-sinf": [
        "Alifbe (2025)", "Ona tili (1-qism)", "Ona tili (2-qism)",
        "O'qish savodxonligi (1-qism)", "O'qish savodxonligi (2-qism)",
        "Matematika (1-qism)", "Matematika (2-qism)", "Matematika (3-qism)",
        "Matematika (4-qism)", "Matematika (2-mashq daftari)",
        "Tabiiy fanlar (1-qism)", "Tabiiy fanlar (2-qism)",
        "Informatika va AT", "Musiqiy savodxonlik", "Tasviriy san'at",
        "Texnologiya", "Texnologiya (mashq daftari)",
        "Tarbiya", "Yozuv daftari", "Jismoniy tarbiya (ish daftari)", "Ingliz tili",
    ],
    "2-sinf": [
        "Ona tili (1-qism)", "Ona tili (2-qism)", "Ona tili (3-qism)", "Ona tili (4-qism)",
        "O'qish savodxonligi (1-qism)", "O'qish savodxonligi (2-qism)",
        "O'qish savodxonligi (3-qism)", "O'qish savodxonligi (4-qism)",
        "O'qish savodxonligi (mashq daftari)",
        "Matematika (1-qism)", "Matematika (2-qism)", "Matematika (3-qism)", "Matematika (4-qism)",
        "Tabiiy fanlar (1-qism)", "Tabiiy fanlar (2-qism)",
        "Informatika va AT", "Musiqiy savodxonlik", "Tasviriy san'at",
        "Texnologiya", "Tarbiya", "Jismoniy tarbiya (ish daftari)",
        "Ingliz tili", "Rus tili",
    ],
    "3-sinf": [
        "Ona tili (1-qism)", "Ona tili (2-qism)", "Ona tili (3-qism)", "Ona tili (4-qism)",
        "O'qish savodxonligi (1-qism)", "O'qish savodxonligi (2-qism)",
        "O'qish savodxonligi (3-qism)", "O'qish savodxonligi (4-qism)",
        "Matematika (1-qism)", "Matematika (2-qism)", "Matematika (3-qism)", "Matematika (4-qism)",
        "Tabiiy fanlar (1-qism)", "Tabiiy fanlar (2-qism)",
        "Musiqiy savodxonlik", "Tasviriy san'at", "Texnologiya", "Tarbiya", "Ingliz tili",
    ],
    "4-sinf": [
        "Ona tili (1-qism)", "Ona tili (2-qism)", "Ona tili (3-qism)", "Ona tili (4-qism)",
        "O'qish savodxonligi (1-qism)", "O'qish savodxonligi (2-qism)",
        "O'qish savodxonligi (3-qism)", "O'qish savodxonligi (4-qism)",
        "Matematika (1-qism)", "Matematika (2-qism)", "Matematika (3-qism)", "Matematika (4-qism)",
        "Tabiiy fanlar (1-qism)", "Tabiiy fanlar (2-qism)", "Tabiiy fanlar (mashq daftari)",
        "Ingliz tili", "Rus tili", "Musiqiy savodxonlik", "Tasviriy san'at",
        "Texnologiya", "Tarbiya", "Jismoniy tarbiya",
    ],
    "5-sinf": [
        "Ona tili", "Adabiyot (1-qism)", "Adabiyot (2-qism)",
        "Matematika (1-qism)", "Matematika (2-qism)",
        "Tarix (Tarixdan hikoyalar)",
        "Tabiiy fanlar (1-qism)", "Tabiiy fanlar (2-qism)",
        "Informatika va AT", "Texnologiya", "Tasviriy san'at",
        "Musiqa madaniyati", "Tarbiya", "Jismoniy tarbiya",
        "Ingliz tili", "Rus tili",
    ],
    "6-sinf": [
        "Ona tili", "Adabiyot (1-qism)", "Adabiyot (2-qism)",
        "Matematika (1-qism)", "Matematika (2-qism)",
        "Tarix (Qadimgi dunyo tarixi)", "Biologiya (Botanika)",
        "Geografiya", "Informatika va AT", "Texnologiya",
        "Tasviriy san'at", "Musiqa madaniyati", "Tarbiya",
        "Jismoniy tarbiya", "Ingliz tili", "Rus tili",
    ],
    "7-sinf": [
        "Ona tili", "Adabiyot (1-qism)", "Adabiyot (2-qism)",
        "Algebra", "Geometriya", "Fizika", "Kimyo",
        "Biologiya (Zoologiya)", "Geografiya",
        "O'zbekiston tarixi", "Jahon tarixi",
        "Informatika va AT", "Texnologiya", "Tasviriy san'at",
        "Musiqa madaniyati", "Tarbiya", "Jismoniy tarbiya",
        "Ingliz tili", "Rus tili",
    ],
    "8-sinf": [
        "Ona tili", "Adabiyot (1-qism)", "Adabiyot (2-qism)",
        "Algebra", "Geometriya", "Fizika", "Kimyo",
        "Biologiya (Odam va uning salomatligi)",
        "Geografiya (O'zbekistonning tabiiy geografiyasi)",
        "O'zbekiston tarixi", "Jahon tarixi",
        "Davlat va huquq asoslari",
        "Informatika va AT", "Chizmachilik", "Tarbiya",
        "Jismoniy tarbiya", "Ingliz tili", "Rus tili",
    ],
    "9-sinf": [
        "Ona tili", "Adabiyot (1-qism)", "Adabiyot (2-qism)",
        "Algebra", "Geometriya", "Fizika", "Kimyo",
        "Biologiya (Sitologiya va genetika asoslari)",
        "Geografiya (O'zbekistonning iqtisodiy geografiyasi)",
        "O'zbekiston tarixi", "Jahon tarixi",
        "Konstitutsiyaviy huquq asoslari",
        "Informatika va AT", "Chizmachilik", "Tarbiya",
        "Jismoniy tarbiya", "CHQBT", "Ingliz tili", "Rus tili",
    ],
    "10-sinf": [
        "Ona tili", "Adabiyot (1-qism)", "Adabiyot (2-qism)",
        "Algebra", "Geometriya", "Fizika", "Kimyo",
        "Biologiya (Umumiy biologiya)",
        "Geografiya (Dunyoning iqtisodiy geografiyasi)",
        "O'zbekiston tarixi", "Jahon tarixi", "Huquqshunoslik",
        "Informatika va AT", "Tarbiya", "Jismoniy tarbiya",
        "CHQBT", "Ingliz tili", "Rus tili", "Tadbirkorlik asoslari",
    ],
    "11-sinf": [
        "Ona tili", "Adabiyot (1-qism)", "Adabiyot (2-qism)",
        "Algebra", "Geometriya", "Fizika", "Kimyo", "Biologiya",
        "Geografiya (Global ekologik muammolar)",
        "O'zbekiston tarixi (Eng yangi tarix)",
        "Jahon tarixi (Eng yangi tarix)", "Inson va jamiyat",
        "Informatika va AT", "Tarbiya", "Jismoniy tarbiya",
        "CHQBT", "Ingliz tili", "Rus tili",
    ],
}

BADIIY_TURLAR = ["Roman", "Hikoya", "She'riyat", "Dunyo adabiyoti", "O'zbek adabiyoti"]
PEDAGOGIK_BOLIMLAR = ["Dars ishlanmalar", "Texnologik xaritalar", "Testlar", "Prezentatsiyalar"]

def make_kb(items, cols=2, back=True):
    rows = []
    for i in range(0, len(items), cols):
        rows.append(items[i:i+cols])
    if back:
        rows.append(["🔙 Orqaga"])
    return ReplyKeyboardMarkup(rows, resize_keyboard=True)

def main_kb():
    return ReplyKeyboardMarkup([
        ["📚 Maktab darsliklari", "📖 Badiiy kitoblar"],
        ["📋 Pedagogik qo'llanmalar"],
        ["🖼️ Rasm tahrirlash", "📄 PDF yasash"],
        ["📝 Word yasash", "💬 Fikr qoldirish"],
    ], resize_keyboard=True)

def back_kb():
    return ReplyKeyboardMarkup([["🔙 Orqaga"]], resize_keyboard=True)

def sinf_kb():
    return ReplyKeyboardMarkup([
        ["1-sinf", "2-sinf", "3-sinf", "4-sinf"],
        ["5-sinf", "6-sinf", "7-sinf", "8-sinf"],
        ["9-sinf", "10-sinf", "11-sinf"],
        ["🔙 Orqaga"]
    ], resize_keyboard=True)

def admin_kb():
    return ReplyKeyboardMarkup([
        ["📚 Darslik yuklash", "📖 Badiiy kitob yuklash"],
        ["📋 Pedagogik yuklash", "🗑 Kitob o'chirish"],
        ["📋 Yuklangan kitoblar", "📨 Xabar yuborish"],
        ["🔙 Chiqish"]
    ], resize_keyboard=True)

def pdf_kb():
    return ReplyKeyboardMarkup([["✅ Tayyor, PDF qil"], ["🔙 Orqaga"]], resize_keyboard=True)

def word_kb():
    return ReplyKeyboardMarkup([["✅ Tayyor, Word qil"], ["🔙 Orqaga"]], resize_keyboard=True)

async def check_sub(user_id, bot):
    try:
        m = await bot.get_chat_member("@pedagoglar1226", user_id)
        if m.status in ["left", "kicked"]:
            return False
    except:
        return False
    return True

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    if user.id == ADMIN_ID:
        context.user_data.clear()
        await update.message.reply_text("👑 Admin paneliga xush kelibsiz!", reply_markup=admin_kb())
        return
    subbed = await check_sub(user.id, context.bot)
    if not subbed:
        kb = InlineKeyboardMarkup([
            [InlineKeyboardButton("➡️ Ilk Bilim", url="https://t.me/+aN43lpwgyEU5NThi")],
            [InlineKeyboardButton("➡️ Pedagoglar", url="https://t.me/pedagoglar1226")],
            [InlineKeyboardButton("✅ Obuna bo'ldim", callback_data="check_sub")],
        ])
        await update.message.reply_text(
            f"👋 Salom {user.first_name}!\n\nBotdan foydalanish uchun kanallarga obuna bo'ling:",
            reply_markup=kb
        )
    else:
        context.user_data.clear()
        await update.message.reply_text(
            f"👋 Salom {user.first_name}! Xush kelibsiz!\n\nBo'limni tanlang:",
            reply_markup=main_kb()
        )

async def callback_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    if query.data == "check_sub":
        subbed = await check_sub(query.from_user.id, context.bot)
        if subbed:
            context.user_data.clear()
            await query.message.reply_text("✅ Rahmat! Xush kelibsiz!", reply_markup=main_kb())
        else:
            await query.answer("❌ Hali obuna bo'lmadingiz!", show_alert=True)

async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    if user.id != ADMIN_ID:
        return
    mode = context.user_data.get("mode", "")
    if mode == "admin_yuklash":
        bolim = context.user_data.get("admin_bolim", "")
        kategoriya = context.user_data.get("admin_kategoriya", "")
        nom = context.user_data.get("kitob_nom", "")
        if not nom:
            await update.message.reply_text("❌ Avval kitob nomini yozing!")
            return
        doc = update.message.document
        if bolim not in file_storage:
            file_storage[bolim] = {}
        if kategoriya not in file_storage[bolim]:
            file_storage[bolim][kategoriya] = {}
        file_storage[bolim][kategoriya][nom] = doc.file_id
        await update.message.reply_text(
            f"✅ '{nom}' saqlandi!\n📂 {bolim} → {kategoriya}",
            reply_markup=admin_kb()
        )
        context.user_data.clear()

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text or ""
    user = update.effective_user
    mode = context.user_data.get("mode", "main")

    # ===== ADMIN =====
    if user.id == ADMIN_ID:
        if text == "🔙 Chiqish" or text == "🔙 Orqaga":
            context.user_data.clear()
            await update.message.reply_text("Admin panel:", reply_markup=admin_kb())
            return

        if text == "📚 Darslik yuklash":
            context.user_data["mode"] = "admin_darslik_sinf"
            context.user_data["admin_bolim"] = "darsliklar"
            await update.message.reply_text("Sinfni tanlang:", reply_markup=sinf_kb())
            return

        if mode == "admin_darslik_sinf" and text in DARSLIKLAR:
            context.user_data["admin_kategoriya"] = text
            context.user_data["mode"] = "admin_darslik_nom"
            kitoblar = DARSLIKLAR[text]
            await update.message.reply_text(
                f"✅ Sinf: {text}\nKitob nomini yozing yoki quyidagilardan tanlang:",
                reply_markup=make_kb(kitoblar)
            )
            return

        if mode == "admin_darslik_nom":
            context.user_data["kitob_nom"] = text
            context.user_data["mode"] = "admin_yuklash"
            await update.message.reply_text(
                f"✅ Kitob: '{text}'\nEndi PDF faylni yuboring:",
                reply_markup=back_kb()
            )
            return

        if text == "📖 Badiiy kitob yuklash":
            context.user_data["mode"] = "admin_badiiy_tur"
            context.user_data["admin_bolim"] = "badiiy"
            await update.message.reply_text("Turni tanlang:", reply_markup=make_kb(BADIIY_TURLAR))
            return

        if mode == "admin_badiiy_tur" and text in BADIIY_TURLAR:
            context.user_data["admin_kategoriya"] = text
            context.user_data["mode"] = "admin_badiiy_nom"
            mav = file_storage["badiiy"].get(text, {})
            mavjud = list(mav.keys())
            msg = f"✅ Tur: {text}\nKitob nomini yozing:"
            if mavjud:
                msg += f"\n\nMavjud kitoblar: {', '.join(mavjud)}"
            await update.message.reply_text(msg, reply_markup=back_kb())
            return

        if mode == "admin_badiiy_nom":
            context.user_data["kitob_nom"] = text
            context.user_data["mode"] = "admin_yuklash"
            await update.message.reply_text(
                f"✅ Kitob: '{text}'\nEndi PDF faylni yuboring:",
                reply_markup=back_kb()
            )
            return

        if text == "📋 Pedagogik yuklash":
            context.user_data["mode"] = "admin_ped_bolim"
            context.user_data["admin_bolim"] = "pedagogik"
            await update.message.reply_text("Bo'limni tanlang:", reply_markup=make_kb(PEDAGOGIK_BOLIMLAR))
            return

        if mode == "admin_ped_bolim" and text in PEDAGOGIK_BOLIMLAR:
            context.user_data["admin_kategoriya"] = text
            context.user_data["mode"] = "admin_ped_nom"
            mav = file_storage["pedagogik"].get(text, {})
            mavjud = list(mav.keys())
            msg = f"✅ Bo'lim: {text}\nQo'llanma nomini yozing:"
            if mavjud:
                msg += f"\n\nMavjud: {', '.join(mavjud)}"
            await update.message.reply_text(msg, reply_markup=back_kb())
            return

        if mode == "admin_ped_nom":
            context.user_data["kitob_nom"] = text
            context.user_data["mode"] = "admin_yuklash"
            await update.message.reply_text(
                f"✅ Nom: '{text}'\nEndi PDF faylni yuboring:",
                reply_markup=back_kb()
            )
            return

        if text == "📋 Yuklangan kitoblar":
            msg = "📋 Yuklangan fayllar:\n\n"
            total = 0
            for bolim, kategoriyalar in file_storage.items():
                for kat, kitoblar in kategoriyalar.items():
                    if kitoblar:
                        msg += f"📂 {bolim} → {kat}:\n"
                        for nom in kitoblar:
                            msg += f"  • {nom}\n"
                            total += 1
            if total == 0:
                msg = "📋 Hali hech narsa yuklanmagan!"
            else:
                msg += f"\nJami: {total} ta fayl"
            await update.message.reply_text(msg, reply_markup=admin_kb())
            return

        if text == "🗑 Kitob o'chirish":
            context.user_data["mode"] = "admin_ochirish"
            msg = "O'chirmoqchi bo'lgan kitob nomini yozing:\n\n"
            for bolim, kategoriyalar in file_storage.items():
                for kat, kitoblar in kategoriyalar.items():
                    for nom in kitoblar:
                        msg += f"• {nom}\n"
            await update.message.reply_text(msg, reply_markup=back_kb())
            return

        if mode == "admin_ochirish":
            deleted = False
            for bolim in file_storage:
                for kat in file_storage[bolim]:
                    if text in file_storage[bolim][kat]:
                        del file_storage[bolim][kat][text]
                        deleted = True
                        break
            if deleted:
                await update.message.reply_text(f"✅ '{text}' o'chirildi!", reply_markup=admin_kb())
            else:
                await update.message.reply_text(f"❌ '{text}' topilmadi!", reply_markup=admin_kb())
            context.user_data.clear()
            return

        if text == "📨 Xabar yuborish":
            context.user_data["mode"] = "admin_xabar"
            await update.message.reply_text("Xabarni yozing:", reply_markup=back_kb())
            return

        if mode == "admin_xabar":
            await update.message.reply_text(f"✅ Xabar yuborildi:\n\n{text}", reply_markup=admin_kb())
            context.user_data.clear()
            return

        return

    # ===== FOYDALANUVCHI =====
    subbed = await check_sub(user.id, context.bot)
    if not subbed:
        await update.message.reply_text("❌ Avval kanallarga obuna bo'ling! /start")
        return

    if text == "🔙 Orqaga":
        if mode == "kitob_list":
            prev = context.user_data.get("prev_mode", "main")
            if prev == "darsliklar":
                context.user_data["mode"] = "darsliklar"
                await update.message.reply_text("📚 Sinfni tanlang:", reply_markup=sinf_kb())
            elif prev == "badiiy":
                context.user_data["mode"] = "badiiy"
                await update.message.reply_text("📖 Turni tanlang:", reply_markup=make_kb(BADIIY_TURLAR))
            elif prev == "pedagogik":
                context.user_data["mode"] = "pedagogik"
                await update.message.reply_text("📋 Bo'limni tanlang:", reply_markup=make_kb(PEDAGOGIK_BOLIMLAR))
            else:
                context.user_data.clear()
                await update.message.reply_text("Asosiy menyu:", reply_markup=main_kb())
        else:
            context.user_data.clear()
            await update.message.reply_text("Asosiy menyu:", reply_markup=main_kb())
        return

    if text == "📚 Maktab darsliklari":
        context.user_data["mode"] = "darsliklar"
        await update.message.reply_text("📚 Sinfni tanlang:", reply_markup=sinf_kb())
        return

    if text == "📖 Badiiy kitoblar":
        context.user_data["mode"] = "badiiy"
        await update.message.reply_text("📖 Turni tanlang:", reply_markup=make_kb(BADIIY_TURLAR))
        return

    if text == "📋 Pedagogik qo'llanmalar":
        context.user_data["mode"] = "pedagogik"
        await update.message.reply_text("📋 Bo'limni tanlang:", reply_markup=make_kb(PEDAGOGIK_BOLIMLAR))
        return

    if text == "🖼️ Rasm tahrirlash":
        context.user_data["mode"] = "tahrirl"
        await update.message.reply_text("🖼️ Rasmni yuboring:", reply_markup=back_kb())
        return

    if text == "📄 PDF yasash":
        context.user_data["mode"] = "pdf"
        context.user_data["pdf_images"] = []
        await update.message.reply_text("📄 Rasmlarni yuboring.\nTugagach '✅ Tayyor, PDF qil' bosing:", reply_markup=pdf_kb())
        return

    if text == "📝 Word yasash":
        context.user_data["mode"] = "word"
        context.user_data["word_items"] = []
        await update.message.reply_text("📝 Rasm yoki matn yuboring.\nTugagach '✅ Tayyor, Word qil' bosing:", reply_markup=word_kb())
        return

    if text == "💬 Fikr qoldirish":
        context.user_data["mode"] = "fikr"
        await update.message.reply_text("💬 Fikringizni yozing:", reply_markup=back_kb())
        return

    if mode == "fikr":
        user_info = f"@{user.username}" if user.username else user.first_name
        await context.bot.send_message(
            chat_id=ADMIN_ID,
            text=f"💬 Yangi fikr!\n👤 {user_info} (ID: {user.id})\n\n📝 {text}"
        )
        await update.message.reply_text("✅ Fikringiz adminga yetkazildi!", reply_markup=main_kb())
        context.user_data.clear()
        return

    # SINF TANLANDI
    if mode == "darsliklar" and text in DARSLIKLAR:
        context.user_data["mode"] = "kitob_list"
        context.user_data["prev_mode"] = "darsliklar"
        context.user_data["kategoriya"] = text
        context.user_data["bolim"] = "darsliklar"
        kitoblar = DARSLIKLAR[text]
        yuklangan = file_storage["darsliklar"].get(text, {})
        rows = []
        for i in range(0, len(kitoblar), 2):
            row = []
            for k in kitoblar[i:i+2]:
                emoji = "✅" if k in yuklangan else "📖"
                row.append(f"{emoji} {k}")
            rows.append(row)
        rows.append(["🔙 Orqaga"])
        await update.message.reply_text(
            f"📚 {text} — kitobni tanlang:\n✅ — yuklab olish mumkin\n📖 — hali yuklanmagan",
            reply_markup=ReplyKeyboardMarkup(rows, resize_keyboard=True)
        )
        return

    # BADIIY TUR TANLANDI
    if mode == "badiiy" and text in BADIIY_TURLAR:
        context.user_data["mode"] = "kitob_list"
        context.user_data["prev_mode"] = "badiiy"
        context.user_data["kategoriya"] = text
        context.user_data["bolim"] = "badiiy"
        kitoblar = list(file_storage["badiiy"].get(text, {}).keys())
        if kitoblar:
            await update.message.reply_text(
                f"📖 {text} — kitobni tanlang:",
                reply_markup=make_kb(kitoblar)
            )
        else:
            await update.message.reply_text(
                f"📖 {text}\n\n⏳ Hali kitob yuklanmagan.\n📥 So'rash: @pedagoglar1226",
                reply_markup=make_kb(BADIIY_TURLAR)
            )
        return

    # PEDAGOGIK BO'LIM TANLANDI
    if mode == "pedagogik" and text in PEDAGOGIK_BOLIMLAR:
        context.user_data["mode"] = "kitob_list"
        context.user_data["prev_mode"] = "pedagogik"
        context.user_data["kategoriya"] = text
        context.user_data["bolim"] = "pedagogik"
        kitoblar = list(file_storage["pedagogik"].get(text, {}).keys())
        if kitoblar:
            await update.message.reply_text(
                f"📋 {text} — tanlang:",
                reply_markup=make_kb(kitoblar)
            )
        else:
            await update.message.reply_text(
                f"📋 {text}\n\n⏳ Hali material yuklanmagan.\n📥 So'rash: @pedagoglar1226",
                reply_markup=make_kb(PEDAGOGIK_BOLIMLAR)
            )
        return

    # KITOB TANLANDI
    if mode == "kitob_list":
        bolim = context.user_data.get("bolim", "")
        kategoriya = context.user_data.get("kategoriya", "")
        clean_text = text.replace("✅ ", "").replace("📖 ", "")
        kitoblar = file_storage.get(bolim, {}).get(kategoriya, {})
        if clean_text in kitoblar:
            await update.message.reply_text(f"📥 '{clean_text}' yuklanmoqda...")
            await context.bot.send_document(
                chat_id=update.effective_chat.id,
                document=kitoblar[clean_text],
                caption=f"📚 {kategoriya} | {clean_text}"
            )
        else:
            await update.message.reply_text(
                f"⏳ '{clean_text}' hali yuklanmagan.\n📥 So'rash: @pedagoglar1226"
            )
        return

    # PDF
    if text == "✅ Tayyor, PDF qil" and mode == "pdf":
        images = context.user_data.get("pdf_images", [])
        if not images:
            await update.message.reply_text("❌ Rasm yuborilmadi!", reply_markup=pdf_kb())
            return
        context.user_data["mode"] = "pdf_nom"
        await update.message.reply_text("📄 PDF uchun nom kiriting:", reply_markup=ReplyKeyboardRemove())
        return

    if mode == "pdf_nom":
        nom = text.strip().replace(" ", "_") or "Hujjat"
        images = context.user_data.get("pdf_images", [])
        await update.message.reply_text("⏳ PDF tayyorlanmoqda...")
        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        pw, ph = A4
        for img_bytes in images:
            img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
            iw, ih = img.size
            ratio = min(pw/iw, ph/ih) * 0.95
            tmp = io.BytesIO()
            img.save(tmp, format="JPEG")
            tmp.seek(0)
            c.drawImage(ImageReader(tmp), (pw-iw*ratio)/2, (ph-ih*ratio)/2, iw*ratio, ih*ratio)
            c.showPage()
        c.save()
        buf.seek(0)
        from telegram import InputFile
        await update.message.reply_document(InputFile(buf, f"{nom}.pdf"), caption=f"✅ {nom}.pdf tayyor!")
        context.user_data.clear()
        await update.message.reply_text("Asosiy menyu:", reply_markup=main_kb())
        return

    # WORD
    if text == "✅ Tayyor, Word qil" and mode == "word":
        items = context.user_data.get("word_items", [])
        if not items:
            await update.message.reply_text("❌ Hech narsa yuborilmadi!", reply_markup=word_kb())
            return
        context.user_data["mode"] = "word_nom"
        await update.message.reply_text("📝 Word uchun nom kiriting:", reply_markup=ReplyKeyboardRemove())
        return

    if mode == "word_nom":
        nom = text.strip().replace(" ", "_") or "Hujjat"
        items = context.user_data.get("word_items", [])
        await update.message.reply_text("⏳ Word tayyorlanmoqda...")
        doc = Document()
        doc.add_heading(nom.replace("_", " "), 0)
        for item in items:
            if item["type"] == "text":
                p = doc.add_paragraph(item["content"])
                p.style.font.size = Pt(12)
            elif item["type"] == "image":
                try:
                    doc.add_picture(io.BytesIO(item["content"]), width=Inches(5))
                except:
                    doc.add_paragraph("[Rasm]")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        from telegram import InputFile
        await update.message.reply_document(InputFile(buf, f"{nom}.docx"), caption=f"✅ {nom}.docx tayyor!")
        context.user_data.clear()
        await update.message.reply_text("Asosiy menyu:", reply_markup=main_kb())
        return

    if mode == "word" and text:
        items = context.user_data.get("word_items", [])
        items.append({"type": "text", "content": text})
        context.user_data["word_items"] = items
        await update.message.reply_text(f"✅ Matn qo'shildi! Jami: {len(items)} ta", reply_markup=word_kb())
        return

async def handle_photo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    subbed = await check_sub(user.id, context.bot)
    if not subbed and user.id != ADMIN_ID:
        await update.message.reply_text("❌ Avval kanallarga obuna bo'ling! /start")
        return

    mode = context.user_data.get("mode", "main")
    photo = update.message.photo[-1]
    file = await context.bot.get_file(photo.file_id)
    file_bytes = bytes(await file.download_as_bytearray())

    if mode == "tahrirl":
        img = Image.open(io.BytesIO(file_bytes)).convert("RGB")
        img = ImageEnhance.Brightness(img).enhance(1.2)
        img = ImageEnhance.Contrast(img).enhance(1.3)
        img = img.filter(ImageFilter.SHARPEN)
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=95)
        buf.seek(0)
        from telegram import InputFile
        await update.message.reply_photo(InputFile(buf, "tahrirlangan.jpg"), caption="✅ Rasm tahrirlandi!")
        return

    if mode == "pdf":
        images = context.user_data.get("pdf_images", [])
        images.append(file_bytes)
        context.user_data["pdf_images"] = images
        await update.message.reply_text(f"✅ Rasm qo'shildi! Jami: {len(images)} ta", reply_markup=pdf_kb())
        return

    if mode == "word":
        items = context.user_data.get("word_items", [])
        items.append({"type": "image", "content": file_bytes})
        context.user_data["word_items"] = items
        await update.message.reply_text(f"✅ Rasm qo'shildi! Jami: {len(items)} ta", reply_markup=word_kb())
        return

    await update.message.reply_text("Asosiy menyudan bo'lim tanlang:", reply_markup=main_kb())

def main():
    app = Application.builder().token(BOT_TOKEN).build()
    app.add_handler(CommandHandler("start", start))
    app.add_handler(CallbackQueryHandler(callback_handler))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    app.add_handler(MessageHandler(filters.PHOTO, handle_photo))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_message))
    print("Bot ishga tushdi!")
    app.run_polling()

if __name__ == "__main__":
    main()
